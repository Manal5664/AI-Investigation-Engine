"""Extractor for DOCX documents.

Extraction walks the WordprocessingML package and reads only paragraphs,
headings, and tables; macros, embedded scripts, and linked content are
never executed. A full-document read is also captured so content is never
lost to paragraph splitting.
"""

from app.documents.base import DocumentExtractionError, DocumentExtractor
from app.documents.models import (
    DocumentKind,
    ExtractedDocument,
    ExtractedPage,
    ExtractedSection,
    UploadedDocument,
)


class DocxExtractor(DocumentExtractor):
    """Extract content from .docx documents via python-docx."""

    def __init__(self) -> None:
        try:
            import docx
        except ImportError as exc:  # pragma: no cover - import guard
            raise DocumentExtractionError(
                "python-docx is not installed; install it to extract DOCX "
                "documents"
            ) from exc
        self._docx = docx

    @property
    def kind(self) -> str:
        return DocumentKind.DOCX.value

    async def extract(
        self,
        uploaded: UploadedDocument,
        content: bytes,
        *,
        max_pages: int,
    ) -> ExtractedDocument:
        try:
            from io import BytesIO

            document = self._docx.Document(BytesIO(content))
        except Exception as exc:
            raise DocumentExtractionError(
                f"document '{uploaded.filename}' could not be parsed as a "
                f"DOCX: {exc}"
            ) from exc

        full_text = self._full_document_text(document)
        page = ExtractedPage(
            page_number=1,
            text=full_text,
            sections=self._collect_sections(document),
        )
        return ExtractedDocument(
            document_id=uploaded.document_id,
            filename=uploaded.filename,
            mime_type=uploaded.mime_type,
            file_size_bytes=uploaded.file_size_bytes,
            content_hash=uploaded.content_hash,
            kind=uploaded.kind,
            extraction_method="docx_paragraphs",
            extracted_at=uploaded.received_at,
            pages=[page],
        )

    def _full_document_text(self, document) -> str:
        parts: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                row_text = " | ".join(c for c in cells if c)
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)

    def _collect_sections(self, document) -> list[ExtractedSection]:
        sections: list[ExtractedSection] = []
        order = 0
        offset = 0
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = (paragraph.style.name or "").lower()
            is_heading = style.startswith("heading") or style == "title"
            sections.append(
                ExtractedSection(
                    heading=text if is_heading else None,
                    text=text if is_heading else text,
                    order_index=order,
                    character_start=offset,
                    character_end=offset + len(text),
                )
            )
            order += 1
            offset += len(text) + 1
        return sections


__all__ = ["DocxExtractor"]
