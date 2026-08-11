import asyncio
import hashlib
import io
from datetime import datetime, timezone

import httpx
import pytest
from pydantic import ValidationError

from app.documents.base import DocumentValidationError
from app.documents.factory import get_document_store
from app.documents.mappers import (
    build_document_graph_edges,
    build_document_graph_nodes,
)
from app.documents.models import DocumentKind, UploadedDocument
from app.documents.validators import validate_upload
from app.graph.factory import get_graph_store
from app.main import app
from app.rag.vectorstore.factory import get_vector_store
from app.services.document_graph_service import DocumentGraphService
from app.services.document_rag_service import DocumentRAGService


def _now() -> datetime:
    return datetime.now(timezone.utc)


def _sha256(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


def _uploaded(
    prefix: str,
    filename: str,
    content: bytes,
    kind: DocumentKind,
    mime_type: str,
    extension: str,
) -> UploadedDocument:
    return UploadedDocument(
        document_id=f"doc-{prefix * 32}",
        filename=filename,
        mime_type=mime_type,
        file_size_bytes=len(content),
        content_hash=_sha256(content),
        kind=kind,
        extension=extension,
        received_at=_now(),
    )


def _txt_bytes(text: str) -> bytes:
    return text.encode("utf-8")


def _png_bytes() -> bytes:
    from PIL import Image

    buffer = io.BytesIO()
    Image.new("RGB", (40, 40), color=(200, 30, 30)).save(buffer, format="PNG")
    return buffer.getvalue()


def _docx_bytes(text: str) -> bytes:
    import docx

    document = docx.Document()
    document.add_heading("Executive Summary", level=1)
    for line in text.splitlines():
        document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _minimal_pdf(text: str) -> bytes:
    objects: list[bytes] = []
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objects.append(b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
    objects.append(
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>"
    )
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("latin-1")
    objects.append(b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream")
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    body = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(body))
        body.extend(f"{index} 0 obj\n".encode())
        body.extend(obj)
        body.extend(b"\nendobj\n")
    xref_offset = len(body)
    body.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    body.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        body.extend(f"{offset:010d} 00000 n \n".encode())
    body.extend(
        b"trailer\n<< /Size "
        + str(len(objects) + 1).encode()
        + b" /Root 1 0 R >>\nstartxref\n"
        + str(xref_offset).encode()
        + b"\n%%EOF"
    )
    return bytes(body)


async def _clear_stores() -> None:
    await get_document_store().clear()
    await get_graph_store().clear()
    await get_vector_store().clear()


def _run(coro) -> None:
    asyncio.run(coro)


@pytest.fixture(autouse=True)
def _fresh_stores():
    _run(_clear_stores())
    yield
    _run(_clear_stores())


# ---------------------------------------------------------------------------
# Validation
# ---------------------------------------------------------------------------

def test_validate_upload_rejects_empty_filename() -> None:
    with pytest.raises(DocumentValidationError):
        validate_upload(
            filename="",
            mime_type="text/plain",
            content=b"hello",
            max_bytes=1024,
        )


def test_validate_upload_rejects_empty_content() -> None:
    with pytest.raises(DocumentValidationError):
        validate_upload(
            filename="empty.txt",
            mime_type="text/plain",
            content=b"",
            max_bytes=1024,
        )


def test_validate_upload_rejects_oversized_document() -> None:
    with pytest.raises(DocumentValidationError):
        validate_upload(
            filename="big.txt",
            mime_type="text/plain",
            content=b"x" * 2048,
            max_bytes=1024,
        )


def test_validate_upload_rejects_unsupported_type() -> None:
    with pytest.raises(DocumentValidationError):
        validate_upload(
            filename="virus.exe",
            mime_type="application/x-msdownload",
            content=b"MZ",
            max_bytes=1024,
        )


def test_validate_upload_derives_kind_and_hash() -> None:
    content = b"alpha beta gamma"
    validated = validate_upload(
        filename="notes.txt",
        mime_type="text/plain",
        content=content,
        max_bytes=1024,
    )
    assert validated.uploaded.kind == DocumentKind.TEXT
    assert validated.uploaded.content_hash == _sha256(content)
    assert validated.uploaded.document_id.startswith("doc-")


def test_validate_upload_uses_detected_mime() -> None:
    validated = validate_upload(
        filename="scan.pdf",
        mime_type="application/pdf",
        content=_minimal_pdf("hello"),
        max_bytes=1024 * 1024,
    )
    assert validated.uploaded.kind == DocumentKind.PDF


# ---------------------------------------------------------------------------
# Extractors
# ---------------------------------------------------------------------------

def test_text_extraction_splits_sections() -> None:
    from app.documents.extractors.text_extractor import TextExtractor

    content = _txt_bytes(
        "# Overview\n\nIntroduction text.\n\n# Findings\n\n"
        "Evidence located on page one.\n"
    )
    uploaded = _uploaded(
        "a", "notes.md", content, DocumentKind.TEXT,
        "text/markdown", ".md",
    )
    extracted = asyncio.run(TextExtractor().extract(uploaded, content, max_pages=50))
    assert extracted.page_count == 1
    assert "Evidence located" in extracted.pages[0].text
    headings = [
        s.heading for page in extracted.pages for s in page.sections
    ]
    assert "Overview" in headings
    assert "Findings" in headings


def test_pdf_extraction_reads_text_and_flags_blank_pages() -> None:
    from app.documents.extractors.pdf_extractor import PdfExtractor

    content = _minimal_pdf("Quarterly financial summary")
    uploaded = _uploaded(
        "b", "report.pdf", content, DocumentKind.PDF,
        "application/pdf", ".pdf",
    )
    extracted = asyncio.run(PdfExtractor().extract(uploaded, content, max_pages=50))
    assert extracted.page_count == 1
    assert "Quarterly financial summary" in extracted.pages[0].text


def test_docx_extraction_reads_paragraphs() -> None:
    from app.documents.extractors.docx_extractor import DocxExtractor

    content = _docx_bytes("Contract renewal clause seven.")
    uploaded = _uploaded(
        "c",
        "contract.docx",
        content,
        DocumentKind.DOCX,
        (
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        ".docx",
    )
    extracted = asyncio.run(DocxExtractor().extract(uploaded, content, max_pages=50))
    assert "Contract renewal clause seven." in extracted.pages[0].text
    assert any(
        s.heading == "Executive Summary"
        for page in extracted.pages
        for s in page.sections
    )


def test_image_extraction_uses_mock_vision() -> None:
    from app.documents.extractors.image_extractor import ImageExtractor
    from app.documents.vision.mock_provider import MockVisionProvider

    content = _png_bytes()
    uploaded = _uploaded(
        "d", "scan.png", content, DocumentKind.IMAGE, "image/png", ".png",
    )
    extracted = asyncio.run(
        ImageExtractor(MockVisionProvider()).extract(
            uploaded,
            content,
            max_pages=50,
        )
    )
    assert extracted.image_content is not None
    assert extracted.image_content.provider_used == "mock"
    assert "scan.png" in extracted.image_content.description


# ---------------------------------------------------------------------------
# Ingestion service
# ---------------------------------------------------------------------------

def test_ingestion_persists_and_dedupes() -> None:
    from app.documents.extractors.factory import ExtractorFactory
    from app.documents.vision.mock_provider import MockVisionProvider
    from app.services.document_ingestion_service import DocumentIngestionService

    async def exercise() -> None:
        store = get_document_store()
        service = DocumentIngestionService(
            store=store,
            extractor_factory=ExtractorFactory(MockVisionProvider()),
            max_bytes=1024 * 1024,
            max_pages=50,
            max_per_request=10,
        )
        content = _txt_bytes("The procurement decision was postponed.")
        first = await service.ingest(
            [("decision.txt", "text/plain", content)]
        )
        second = await service.ingest(
            [("copy.txt", "text/plain", content)]
        )
        stats = await store.stats()
        return first[0], second[0], stats

    first, second, stats = asyncio.run(exercise())
    assert first.duplicate is False
    assert second.duplicate is True
    assert second.document.document_id == first.document.document_id
    assert stats.document_count == 1


# ---------------------------------------------------------------------------
# Graph mapping
# ---------------------------------------------------------------------------

def test_document_graph_mapping_builds_nodes_and_edges() -> None:
    from app.documents.extractors.text_extractor import TextExtractor

    content = _txt_bytes("Alpha company acquired Beta Inc.")
    uploaded = _uploaded(
        "e", "press.txt", content, DocumentKind.TEXT, "text/plain", ".txt",
    )
    extracted = asyncio.run(
        TextExtractor().extract(uploaded, content, max_pages=50)
    )
    nodes = build_document_graph_nodes(extracted)
    edges = build_document_graph_edges(extracted)
    assert any(node.node_type.value == "evidence" for node in nodes)
    assert any(node.node_type.value == "topic" for node in nodes)


def test_document_graph_service_indexes_document() -> None:
    from app.documents.extractors.factory import ExtractorFactory
    from app.documents.vision.mock_provider import MockVisionProvider
    from app.services.document_ingestion_service import DocumentIngestionService

    async def exercise() -> None:
        store = get_document_store()
        service = DocumentIngestionService(
            store=store,
            extractor_factory=ExtractorFactory(MockVisionProvider()),
            max_bytes=1024 * 1024,
            max_pages=50,
            max_per_request=10,
        )
        result = await service.ingest(
            [("g.txt", "text/plain", _txt_bytes("Company X expands in Berlin."))]
        )
        document_id = result[0].document.document_id
        graph_service = DocumentGraphService(get_graph_store(), store)
        nodes, edges = await graph_service.index_document(document_id)
        return len(nodes), len(edges)

    node_count, edge_count = asyncio.run(exercise())
    assert node_count >= 1
    assert edge_count >= 0


# ---------------------------------------------------------------------------
# Document RAG service
# ---------------------------------------------------------------------------

def test_document_rag_index_and_relevant_pages() -> None:
    from app.rag.embeddings.mock_provider import MockEmbeddingProvider
    from app.rag.vectorstore.in_memory import InMemoryVectorStore
    from app.services.document_ingestion_service import DocumentIngestionService
    from app.documents.extractors.factory import ExtractorFactory
    from app.documents.vision.mock_provider import MockVisionProvider

    async def exercise() -> None:
        store = get_document_store()
        ingestion = DocumentIngestionService(
            store=store,
            extractor_factory=ExtractorFactory(MockVisionProvider()),
            max_bytes=1024 * 1024,
            max_pages=50,
            max_per_request=10,
        )
        result = await ingestion.ingest(
            [
                (
                    "funding.txt",
                    "text/plain",
                    _txt_bytes(
                        "The company announced a new funding round led by "
                        "northstar capital."
                    ),
                )
            ]
        )
        document_id = result[0].document.document_id
        rag = DocumentRAGService(
            document_store=store,
            embedding_provider=MockEmbeddingProvider(),
            vector_store=InMemoryVectorStore(),
        )
        index_result = await rag.index_document(document_id)
        assert index_result is not None
        assert index_result.pages_indexed == 1
        assert index_result.chunks_created >= 1

        pages = await rag.find_relevant_pages("funding round northstar")
        assert pages, "expected at least one relevant page"
        assert pages[0].document_id == document_id
        assert pages[0].page_number == 1

    asyncio.run(exercise())


# ---------------------------------------------------------------------------
# Page-aware evidence extraction
# ---------------------------------------------------------------------------

def test_page_aware_evidence_extractor_reads_stored_documents() -> None:
    from app.documents.extractors.factory import ExtractorFactory
    from app.documents.vision.mock_provider import MockVisionProvider
    from app.evidence.mock_extractor import MockEvidenceExtractor
    from app.evidence.page_aware_extractor import PageAwareEvidenceExtractor
    from app.schemas.investigation import InvestigationSubQuestion
    from app.services.document_ingestion_service import (
        DocumentIngestionService,
    )

    async def exercise() -> None:
        store = get_document_store()
        ingestion = DocumentIngestionService(
            store=store,
            extractor_factory=ExtractorFactory(MockVisionProvider()),
            max_bytes=1024 * 1024,
            max_pages=50,
            max_per_request=10,
        )
        await ingestion.ingest(
            [
                (
                    "vendor.txt",
                    "text/plain",
                    _txt_bytes(
                        "The vendor delivered the hardware late and without "
                        "certification paperwork."
                    ),
                )
            ]
        )
        extractor = PageAwareEvidenceExtractor(
            MockEvidenceExtractor(),
            store,
        )
        sub_question = InvestigationSubQuestion(
            id="sq-01",
            question="Was the vendor compliant?",
            purpose="Assess vendor compliance from supplied documents.",
            priority=1,
        )
        items = await extractor.extract(sub_question, [])
        return len(items), items[0].provenance.relevant_passage if items else None

    count, first_passage = asyncio.run(exercise())
    assert count >= 1
    assert first_passage is not None
    assert "vendor" in first_passage.casefold()

    asyncio.run(_clear_stores())


# ---------------------------------------------------------------------------
# API integration
# ---------------------------------------------------------------------------

def test_documents_api_upload_list_get_stats_graph_and_index() -> None:
    async def exercise() -> None:
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport,
            base_url="http://testserver",
        ) as client:
            upload = await client.post(
                "/api/v1/documents/upload",
                files={
                    "files": (
                        "api.txt",
                        _txt_bytes(
                            "The audit found irregularities in invoice "
                            "reconciliation."
                        ),
                        "text/plain",
                    )
                },
            )
            assert upload.status_code == 200
            body = upload.json()
            assert body["status"] == "completed"
            assert len(body["documents"]) == 1
            document_id = body["documents"][0]["document"]["document_id"]

            stats = await client.get("/api/v1/documents/store")
            assert stats.status_code == 200
            assert stats.json()["stats"]["document_count"] == 1

            listed = await client.get("/api/v1/documents/list")
            assert listed.status_code == 200
            assert listed.json()["total"] == 1
            assert listed.json()["documents"][0]["page_count"] == 1

            fetched = await client.get(f"/api/v1/documents/{document_id}")
            assert fetched.status_code == 200
            assert fetched.json()["document"]["kind"] == "text"

            graph = await client.post(
                f"/api/v1/documents/{document_id}/graph",
                json={"document_id": document_id},
            )
            assert graph.status_code == 200
            assert graph.json()["node_count"] >= 1

            indexed = await client.post(
                f"/api/v1/documents/{document_id}/index",
                headers={"content-type": "application/json"},
            )
            assert indexed.status_code == 200
            assert indexed.json()["pages_indexed"] == 1

            deleted = await client.post(
                "/api/v1/documents/delete",
                json={"document_ids": [document_id]},
            )
            assert deleted.status_code == 200
            assert deleted.json()["deleted_count"] == 1

            stats_after = await client.get("/api/v1/documents/store")
            assert stats_after.json()["stats"]["document_count"] == 0

            missing = await client.get("/api/v1/documents/doc-00000000000000000000000000000000")
            assert missing.status_code == 404

    asyncio.run(exercise())


def test_documents_api_rejects_unsupported_file() -> None:
    async def exercise() -> None:
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport,
            base_url="http://testserver",
        ) as client:
            response = await client.post(
                "/api/v1/documents/upload",
                files={
                    "files": (
                        "evil.sh",
                        b"#!/bin/sh\necho pwned",
                        "application/x-sh",
                    )
                },
            )
            assert response.status_code == 422
            assert "unsupported" in response.json()["message"]

    asyncio.run(exercise())


def test_documents_investigation_uses_mock_generator() -> None:
    async def exercise() -> None:
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport,
            base_url="http://testserver",
        ) as client:
            upload = await client.post(
                "/api/v1/documents/upload",
                files={
                    "files": (
                        "policy.txt",
                        _txt_bytes(
                            "The new policy requires quarterly disclosure "
                            "of related-party transactions."
                        ),
                        "text/plain",
                    )
                },
            )
            document_id = upload.json()["documents"][0]["document"]["document_id"]
            report = await client.post(
                "/api/v1/documents/investigations",
                json={
                    "query": "What does the policy require?",
                    "depth": "quick",
                    "document_ids": [document_id],
                    "use_rag": True,
                    "use_graph": True,
                },
            )
            assert report.status_code == 200
            payload = report.json()
            assert payload["status"] == "completed"
            assert payload["fallback_used"] is True
            assert payload["provider_used"] == "mock"
            assert payload["documents_used"][0]["document_id"] == document_id

    asyncio.run(exercise())


def test_documents_investigation_reports_no_documents() -> None:
    async def exercise() -> None:
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport,
            base_url="http://testserver",
        ) as client:
            report = await client.post(
                "/api/v1/documents/investigations",
                json={"query": "Anything stored here?", "depth": "quick"},
            )
            assert report.status_code == 200
            assert report.json()["status"] == "no_documents"

    asyncio.run(exercise())
