"""Run a real end-to-end document experiment against the ingestion service.

Uploads a generated multi-page text/PDF document, indexes it into the graph
and vector stores, then runs a document-grounded investigation.
"""

import asyncio
import io

from app.core.config import settings
from app.documents.extractors.factory import ExtractorFactory
from app.documents.factory import get_document_store
from app.documents.reporting_factory import create_document_report_generator
from app.documents.vision.factory import create_vision_provider
from app.graph.factory import get_graph_store
from app.rag.embeddings.factory import create_embedding_provider
from app.rag.vectorstore.factory import get_vector_store
from app.schemas.documents import DocumentInvestigationRequest
from app.services.document_graph_service import DocumentGraphService
from app.services.document_ingestion_service import DocumentIngestionService
from app.services.document_investigation_service import (
    DocumentInvestigationService,
)
from app.services.document_rag_service import DocumentRAGService

DOCUMENT_TEXT = """\
# Executive Summary
The procurement committee postponed the contract award pending a review of
the two remaining bidders.

# Findings
Bidder A provided audited financial statements but missed the compliance
deadline. Bidder B submitted on time but flagged a conflict of interest.

# Recommendation
The committee should re-evaluate both bids and publish the reasons for the
final award.
"""


def _docx_bytes(text: str) -> bytes:
    import docx

    document = docx.Document()
    for line in text.splitlines():
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.strip():
            document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def main() -> int:
    store = get_document_store()
    ingestion = DocumentIngestionService(
        store=store,
        extractor_factory=ExtractorFactory(create_vision_provider(settings)),
        max_bytes=10 * 1024 * 1024,
        max_pages=50,
        max_per_request=10,
    )

    results = await ingestion.ingest(
        [
            (
                "committee.docx",
                (
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                _docx_bytes(DOCUMENT_TEXT),
            ),
        ]
    )
    for result in results:
        uploaded = result.document
        print(
            f"[ingested] {uploaded.filename} -> {uploaded.document_id} "
            f"({result.page_count} pages, {result.character_count} chars, "
            f"duplicate={result.duplicate})"
        )

    document_id = results[0].document.document_id

    graph_service = DocumentGraphService(get_graph_store(), store)
    nodes, edges = await graph_service.index_document(document_id)
    print(f"[graph] {len(nodes)} nodes, {len(edges)} edges indexed")

    rag_service = DocumentRAGService(
        document_store=store,
        embedding_provider=create_embedding_provider(),
        vector_store=get_vector_store(),
    )
    index_result = await rag_service.index_document(document_id)
    print(
        f"[rag] {index_result.pages_indexed} pages indexed, "
        f"{index_result.chunks_created} chunks created"
    )

    report = await DocumentInvestigationService(
        document_store=store,
        generator=create_document_report_generator(),
        rag_service=rag_service,
        graph_store=get_graph_store(),
    ).investigate(
        DocumentInvestigationRequest(
            query=(
                "What compliance problems did the bidders face during the "
                "procurement?"
            ),
            depth="standard",
            document_ids=[document_id],
        )
    )
    print(f"[report] status={report.status} fallback={report.fallback_used}")
    for finding in report.findings:
        print(f"  - {finding}")
    return 0


if __name__ == "__main__":
    raise SystemExit(asyncio.run(main()))
